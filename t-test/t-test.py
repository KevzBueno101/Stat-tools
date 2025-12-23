import customtkinter as ctk
from tkinter import filedialog, messagebox
from scipy import stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime

class IndependentTTestApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Independent Samples t-test")
        self.root.geometry("900x800")
        
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        
        # Results storage
        self.results = None
        
        # Main frame
        main_frame = ctk.CTkFrame(self.root)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        title_label = ctk.CTkLabel(
            main_frame, 
            text="Independent Samples t-test", 
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(0, 20))
        
        # Document metadata frame
        metadata_frame = ctk.CTkFrame(main_frame)
        metadata_frame.pack(fill="x", padx=10, pady=(0, 15))
        
        # Report title entry
        report_title_label = ctk.CTkLabel(
            metadata_frame, 
            text="Report Title:", 
            font=ctk.CTkFont(size=13, weight="bold")
        )
        report_title_label.grid(row=0, column=0, sticky="w", padx=10, pady=5)
        
        self.report_title_entry = ctk.CTkEntry(
            metadata_frame, 
            width=500,
            placeholder_text="Enter report title (optional)"
        )
        self.report_title_entry.grid(row=0, column=1, padx=10, pady=5)
        
        # Byline entry
        byline_label = ctk.CTkLabel(
            metadata_frame, 
            text="By:", 
            font=ctk.CTkFont(size=13, weight="bold")
        )
        byline_label.grid(row=1, column=0, sticky="w", padx=10, pady=5)
        
        self.byline_entry = ctk.CTkEntry(
            metadata_frame, 
            width=500,
            placeholder_text="Enter author name (optional)"
        )
        self.byline_entry.grid(row=1, column=1, padx=10, pady=5)
        
        # Separator
        separator = ctk.CTkFrame(main_frame, height=2, fg_color="gray70")
        separator.pack(fill="x", padx=10, pady=10)
        
        # Group 1 Input
        group1_label = ctk.CTkLabel(
            main_frame, 
            text="Group 1 Scores (comma or newline separated):", 
            font=ctk.CTkFont(size=14, weight="bold")
        )
        group1_label.pack(anchor="w", padx=10)
        
        self.group1_text = ctk.CTkTextbox(main_frame, height=100, width=700)
        self.group1_text.pack(padx=10, pady=(5, 15))
        
        # Group 2 Input
        group2_label = ctk.CTkLabel(
            main_frame, 
            text="Group 2 Scores (comma or newline separated):", 
            font=ctk.CTkFont(size=14, weight="bold")
        )
        group2_label.pack(anchor="w", padx=10)
        
        self.group2_text = ctk.CTkTextbox(main_frame, height=100, width=700)
        self.group2_text.pack(padx=10, pady=(5, 15))
        
        # Button frame
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(pady=10)
        
        # Compute button
        self.compute_button = ctk.CTkButton(
            button_frame, 
            text="Compute t-test", 
            command=self.compute_ttest,
            font=ctk.CTkFont(size=14, weight="bold"),
            width=150,
            height=35
        )
        self.compute_button.pack(side="left", padx=5)
        
        # Clear button
        self.clear_button = ctk.CTkButton(
            button_frame, 
            text="Clear", 
            command=self.clear_fields,
            font=ctk.CTkFont(size=14),
            width=100,
            height=35
        )
        self.clear_button.pack(side="left", padx=5)
        
        # Save button
        self.save_button = ctk.CTkButton(
            button_frame, 
            text="Save Results to DOCX", 
            command=self.save_to_docx,
            font=ctk.CTkFont(size=14, weight="bold"),
            width=180,
            height=35,
            state="disabled"
        )
        self.save_button.pack(side="left", padx=5)
        
        # Status label
        self.status_label = ctk.CTkLabel(
            main_frame, 
            text="", 
            font=ctk.CTkFont(size=12)
        )
        self.status_label.pack(pady=5)
        
        # Results display
        results_label = ctk.CTkLabel(
            main_frame, 
            text="Results:", 
            font=ctk.CTkFont(size=14, weight="bold")
        )
        results_label.pack(anchor="w", padx=10, pady=(10, 5))
        
        self.results_text = ctk.CTkTextbox(main_frame, height=180, width=700)
        self.results_text.pack(padx=10, pady=(0, 10))
        self.results_text.configure(state="disabled")
    
    def parse_input(self, text):
        """Parse input text and extract numeric values"""
        # Replace newlines and commas with spaces
        text = re.sub(r'[,\n\r\t]+', ' ', text)
        # Extract all numeric values (including decimals and negatives)
        values = re.findall(r'-?\d+\.?\d*', text)
        # Convert to float
        return [float(v) for v in values if v]
    
    def compute_ttest(self):
        """Compute independent samples t-test"""
        try:
            # Get input data
            group1_text = self.group1_text.get("1.0", "end").strip()
            group2_text = self.group2_text.get("1.0", "end").strip()
            
            # Validate inputs are not empty
            if not group1_text or not group2_text:
                messagebox.showerror("Input Error", "Please enter data for both groups.")
                return
            
            # Parse inputs
            group1_data = self.parse_input(group1_text)
            group2_data = self.parse_input(group2_text)
            
            # Validate we have data
            if len(group1_data) == 0 or len(group2_data) == 0:
                messagebox.showerror("Input Error", "No valid numeric values found. Please enter numeric data.")
                return
            
            # Validate minimum sample size
            if len(group1_data) < 2 or len(group2_data) < 2:
                messagebox.showerror("Input Error", "Each group must have at least 2 values.")
                return
            
            # Compute statistics
            mean1 = sum(group1_data) / len(group1_data)
            mean2 = sum(group2_data) / len(group2_data)
            n1 = len(group1_data)
            n2 = len(group2_data)
            
            # Perform Welch's t-test (equal_var=False)
            t_statistic, p_value = stats.ttest_ind(group1_data, group2_data, equal_var=False)
            
            # Calculate degrees of freedom for Welch's t-test
            var1 = sum((x - mean1) ** 2 for x in group1_data) / (n1 - 1)
            var2 = sum((x - mean2) ** 2 for x in group2_data) / (n2 - 1)
            df = ((var1/n1 + var2/n2) ** 2) / ((var1/n1)**2/(n1-1) + (var2/n2)**2/(n2-1))
            
            # Determine significance
            alpha = 0.05
            if p_value <= alpha:
                conclusion = "There is a significant difference between the two groups."
            else:
                conclusion = "There is no significant difference between the two groups."
            
            # Store results
            self.results = {
                'group1_data': group1_data,
                'group2_data': group2_data,
                'mean1': mean1,
                'mean2': mean2,
                'n1': n1,
                'n2': n2,
                't_statistic': t_statistic,
                'df': df,
                'p_value': p_value,
                'conclusion': conclusion,
                'timestamp': datetime.now()
            }
            
            # Display results
            results_text = f"""
Independent Samples t-test Results
{'=' * 50}

Group 1:
  Mean: {mean1:.4f}
  Sample Size: {n1}

Group 2:
  Mean: {mean2:.4f}
  Sample Size: {n2}

Statistical Results:
  t({df:.2f}) = {t_statistic:.4f}
  p-value (two-tailed) = {p_value:.4f}
  Alpha level: {alpha}

Conclusion:
  {conclusion}
"""
            
            self.results_text.configure(state="normal")
            self.results_text.delete("1.0", "end")
            self.results_text.insert("1.0", results_text)
            self.results_text.configure(state="disabled")
            
            # Update status
            self.status_label.configure(text="✓ Computation successful", text_color="green")
            
            # Enable save button
            self.save_button.configure(state="normal")
            
        except ValueError as e:
            messagebox.showerror("Input Error", f"Invalid input: {str(e)}\nPlease enter only numeric values.")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
    
    def save_to_docx(self):
        """Save results to a DOCX file"""
        if self.results is None:
            messagebox.showwarning("No Results", "Please compute the t-test first.")
            return
        
        try:
            # Ask user for save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Save Results"
            )
            
            if not file_path:
                return
            
            # Get metadata
            report_title = self.report_title_entry.get().strip()
            byline = self.byline_entry.get().strip()
            
            # Use default title if none provided
            if not report_title:
                report_title = "Independent Samples t-test Results"
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(report_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add byline if provided
            if byline:
                byline_paragraph = doc.add_paragraph(f"By: {byline}")
                byline_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                byline_run = byline_paragraph.runs[0]
                byline_run.font.size = Pt(12)
                byline_run.italic = True
            
            # Add timestamp
            timestamp_str = self.results['timestamp'].strftime("%B %d, %Y at %I:%M %p")
            timestamp_paragraph = doc.add_paragraph(f"Generated: {timestamp_str}")
            timestamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_run = timestamp_paragraph.runs[0]
            timestamp_run.font.size = Pt(10)
            timestamp_run.font.color.rgb = None  # Gray color
            
            doc.add_paragraph()
            
            # Add descriptive statistics table
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Group'
            header_cells[1].text = 'Mean'
            header_cells[2].text = 'Sample Size'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Group 1 data
            row1_cells = table.rows[1].cells
            row1_cells[0].text = 'Group 1'
            row1_cells[1].text = f"{self.results['mean1']:.4f}"
            row1_cells[2].text = str(self.results['n1'])
            
            # Group 2 data
            row2_cells = table.rows[2].cells
            row2_cells[0].text = 'Group 2'
            row2_cells[1].text = f"{self.results['mean2']:.4f}"
            row2_cells[2].text = str(self.results['n2'])
            
            doc.add_paragraph()
            
            # Add statistical results
            results_heading = doc.add_heading('Statistical Results', 2)
            
            stats_text = (
                f"t({self.results['df']:.2f}) = {self.results['t_statistic']:.4f}, "
                f"p = {self.results['p_value']:.4f}"
            )
            doc.add_paragraph(stats_text)
            
            doc.add_paragraph()
            
            # Add conclusion
            conclusion_heading = doc.add_heading('Conclusion', 2)
            doc.add_paragraph(self.results['conclusion'])
            
            # Save document
            doc.save(file_path)
            
            # Update status
            self.status_label.configure(text=f"✓ Results saved to {file_path}", text_color="green")
            messagebox.showinfo("Success", f"Results saved successfully to:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save file: {str(e)}")
    
    def clear_fields(self):
        """Clear all input and output fields"""
        self.group1_text.delete("1.0", "end")
        self.group2_text.delete("1.0", "end")
        self.results_text.configure(state="normal")
        self.results_text.delete("1.0", "end")
        self.results_text.configure(state="disabled")
        self.report_title_entry.delete(0, "end")
        self.byline_entry.delete(0, "end")
        self.status_label.configure(text="")
        self.results = None
        self.save_button.configure(state="disabled")

def main():
    root = ctk.CTk()
    app = IndependentTTestApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()