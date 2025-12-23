"""
Main Application File
Launches the CustomTkinter GUI for Pearson's r Critical Value Calculator.
"""

import customtkinter as ctk
from tkinter import messagebox
import threading
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import custom modules
from stats_utils import (
    validate_sample_size,
    validate_alpha,
    compute_pearson_r_critical,
    format_results
)
from gui_components import LabeledEntry, ResultsDisplay, ExcelAnalysisFrame
from excel_utils import analyze_excel_correlation


class PearsonCalculatorApp(ctk.CTk):
    """
    Main application class for Pearson's r Critical Value Calculator.
    """
    
    def __init__(self):
        """Initialize the application."""
        super().__init__()
        
        # Configure window
        self.title("Pearson's r Critical Value Calculator")
        self.geometry("1200x700")  # Larger window for tabbed interface
        
        # Center window on screen
        self.center_window()
        
        # Set appearance
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # Configure grid for main window
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        # Store current results for saving
        self.current_results = None
        
        # Create GUI elements
        self.create_widgets()
    
    def center_window(self):
        """Center the application window on the screen."""
        self.update_idletasks()
        
        # Get screen dimensions
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        # Get window dimensions
        window_width = 1200
        window_height = 700
        
        # Calculate position
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        # Set position
        self.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    def create_widgets(self):
        """Create and layout all GUI widgets with tabbed interface."""
        
        # ===== MAIN CONTAINER =====
        main_frame = ctk.CTkFrame(self)
        main_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(1, weight=1)
        
        # ===== HEADER =====
        header_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=20, pady=(10, 5), sticky="ew")
        
        title_label = ctk.CTkLabel(
            header_frame,
            text="Pearson's r Critical Value Calculator",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(0, 5))
        
        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Manual Calculation & Automated Excel Analysis",
            font=ctk.CTkFont(size=13),
            text_color="gray50"
        )
        subtitle_label.pack()
        
        # ===== TABVIEW =====
        self.tabview = ctk.CTkTabview(main_frame)
        self.tabview.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        
        # Create tabs
        self.tabview.add("Manual Calculation")
        self.tabview.add("Excel Analysis")
        
        # Configure tab grids
        self.tabview.tab("Manual Calculation").grid_columnconfigure(0, weight=1)
        self.tabview.tab("Manual Calculation").grid_columnconfigure(1, weight=2)
        self.tabview.tab("Manual Calculation").grid_rowconfigure(0, weight=1)
        
        self.tabview.tab("Excel Analysis").grid_columnconfigure(0, weight=1)
        self.tabview.tab("Excel Analysis").grid_columnconfigure(1, weight=2)
        self.tabview.tab("Excel Analysis").grid_rowconfigure(0, weight=1)
        
        # ===== TAB 1: MANUAL CALCULATION =====
        self.create_manual_tab()
        
        # ===== TAB 2: EXCEL ANALYSIS =====
        self.create_excel_tab()
    
    def create_manual_tab(self):
        """Create the manual calculation tab."""
        tab = self.tabview.tab("Manual Calculation")
        
        # Left column
        left_frame = ctk.CTkFrame(tab)
        left_frame.grid(row=0, column=0, padx=(10, 5), pady=10, sticky="nsew")
        left_frame.grid_columnconfigure(0, weight=1)
        left_frame.grid_rowconfigure(3, weight=1)
        
        # Input section
        input_frame = ctk.CTkFrame(left_frame)
        input_frame.grid(row=0, column=0, padx=15, pady=(15, 10), sticky="ew")
        input_frame.grid_columnconfigure(0, weight=1)
        
        # Sample size input
        self.sample_size_entry = LabeledEntry(
            input_frame,
            label_text="Sample Size (n):",
            default_value="30"
        )
        self.sample_size_entry.grid(row=0, column=0, padx=15, pady=(15, 5), sticky="ew")
        
        # Alpha input
        self.alpha_entry = LabeledEntry(
            input_frame,
            label_text="Significance Level (α):",
            default_value="0.05"
        )
        self.alpha_entry.grid(row=1, column=0, padx=15, pady=5, sticky="ew")
        
        # Test type selection
        test_type_frame = ctk.CTkFrame(input_frame, fg_color="transparent")
        test_type_frame.grid(row=2, column=0, padx=15, pady=(5, 15), sticky="ew")
        
        test_type_label = ctk.CTkLabel(
            test_type_frame,
            text="Test Type:",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        test_type_label.pack(anchor="w", padx=5, pady=(0, 5))
        
        self.test_type_var = ctk.StringVar(value="two-tailed")
        
        test_type_menu = ctk.CTkOptionMenu(
            test_type_frame,
            variable=self.test_type_var,
            values=["two-tailed", "one-tailed"],
            font=ctk.CTkFont(size=13)
        )
        test_type_menu.pack(fill="x", padx=5)
        
        # Button section
        button_frame = ctk.CTkFrame(left_frame, fg_color="transparent")
        button_frame.grid(row=1, column=0, padx=15, pady=10, sticky="ew")
        button_frame.grid_columnconfigure((0, 1), weight=1)
        
        # Calculate button
        self.calculate_button = ctk.CTkButton(
            button_frame,
            text="Compute Critical Value",
            command=self.on_calculate_click,
            font=ctk.CTkFont(size=14, weight="bold"),
            height=40
        )
        self.calculate_button.grid(row=0, column=0, padx=(0, 5), sticky="ew")
        
        # Save button
        self.save_button = ctk.CTkButton(
            button_frame,
            text="Save Results",
            command=self.save_results,
            font=ctk.CTkFont(size=14, weight="bold"),
            height=40,
            fg_color="gray60",
            hover_color="gray50"
        )
        self.save_button.grid(row=0, column=1, padx=(5, 0), sticky="ew")
        
        # Footer section
        footer_frame = ctk.CTkFrame(left_frame, fg_color="transparent")
        footer_frame.grid(row=2, column=0, padx=15, pady=(10, 15), sticky="s")
        
        formula_label = ctk.CTkLabel(
            footer_frame,
            text="Formula: r = √(t² / (t² + df))  |  df = n - 2",
            font=ctk.CTkFont(size=11),
            text_color="gray50"
        )
        formula_label.pack()
        
        # Right column (Results)
        right_frame = ctk.CTkFrame(tab)
        right_frame.grid(row=0, column=1, padx=(5, 10), pady=10, sticky="nsew")
        right_frame.grid_columnconfigure(0, weight=1)
        right_frame.grid_rowconfigure(0, weight=1)
        
        # Results display
        self.results_display = ResultsDisplay(right_frame)
        self.results_display.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
    
    def create_excel_tab(self):
        """Create the Excel analysis tab."""
        tab = self.tabview.tab("Excel Analysis")
        
        # Left column - Excel controls
        left_frame = ctk.CTkFrame(tab)
        left_frame.grid(row=0, column=0, padx=(10, 5), pady=10, sticky="nsew")
        left_frame.grid_columnconfigure(0, weight=1)
        left_frame.grid_rowconfigure(2, weight=1)
        
        # Excel analysis frame
        self.excel_frame = ExcelAnalysisFrame(
            left_frame,
            on_analyze_callback=self.analyze_excel_file
        )
        self.excel_frame.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="ew")
        
        # Analysis parameters frame
        params_frame = ctk.CTkFrame(left_frame)
        params_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        params_frame.grid_columnconfigure(0, weight=1)
        
        params_title = ctk.CTkLabel(
            params_frame,
            text="Analysis Parameters",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        params_title.grid(row=0, column=0, padx=15, pady=(15, 10), sticky="w")
        
        # Alpha input for Excel
        self.excel_alpha_entry = LabeledEntry(
            params_frame,
            label_text="Significance Level (α):",
            default_value="0.05"
        )
        self.excel_alpha_entry.grid(row=1, column=0, padx=15, pady=5, sticky="ew")
        
        # Test type for Excel
        test_type_frame = ctk.CTkFrame(params_frame, fg_color="transparent")
        test_type_frame.grid(row=2, column=0, padx=15, pady=(5, 15), sticky="ew")
        
        test_type_label = ctk.CTkLabel(
            test_type_frame,
            text="Test Type:",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        test_type_label.pack(anchor="w", padx=5, pady=(0, 5))
        
        self.excel_test_type_var = ctk.StringVar(value="two-tailed")
        
        test_type_menu = ctk.CTkOptionMenu(
            test_type_frame,
            variable=self.excel_test_type_var,
            values=["two-tailed", "one-tailed"],
            font=ctk.CTkFont(size=13)
        )
        test_type_menu.pack(fill="x", padx=5)
        
        # Right column (Results)
        right_frame = ctk.CTkFrame(tab)
        right_frame.grid(row=0, column=1, padx=(5, 10), pady=10, sticky="nsew")
        right_frame.grid_columnconfigure(0, weight=1)
        right_frame.grid_rowconfigure(0, weight=1)
        
        # Results display for Excel
        self.excel_results_display = ResultsDisplay(right_frame)
        self.excel_results_display.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
    
    def validate_inputs(self) -> bool:
        """
        Validate all input fields.
        
        Returns:
        --------
        bool : True if all inputs are valid, False otherwise
        """
        # Validate sample size
        is_valid, error_msg = validate_sample_size(self.sample_size_entry.get())
        if not is_valid:
            messagebox.showerror("Input Error", error_msg)
            return False
        
        # Validate alpha
        is_valid, error_msg = validate_alpha(self.alpha_entry.get())
        if not is_valid:
            messagebox.showerror("Input Error", error_msg)
            return False
        
        return True
    
    def calculate_in_thread(self):
        """
        Perform calculation in a separate thread to avoid UI freezing.
        This ensures the GUI remains responsive during computation.
        """
        try:
            # Get input values
            n = int(float(self.sample_size_entry.get()))
            alpha = float(self.alpha_entry.get())
            test_type = self.test_type_var.get()
            
            # Compute results
            results = compute_pearson_r_critical(n, alpha, test_type)
            
            # Store results for saving
            self.current_results = results
            
            # Update GUI in main thread
            self.after(0, lambda: self.display_results(results))
            
        except Exception as e:
            # Show error in main thread
            self.after(0, lambda: messagebox.showerror(
                "Calculation Error",
                f"An error occurred during calculation:\n{str(e)}"
            ))
        finally:
            # Re-enable button in main thread
            self.after(0, lambda: self.calculate_button.configure(state="normal"))
    
    def on_calculate_click(self):
        """Handle calculate button click event."""
        # Validate inputs first
        if not self.validate_inputs():
            return
        
        # Disable button to prevent multiple clicks
        self.calculate_button.configure(state="disabled")
        
        # Run calculation in separate thread
        calc_thread = threading.Thread(target=self.calculate_in_thread)
        calc_thread.daemon = True
        calc_thread.start()
    
    def display_results(self, results: dict):
        """
        Display calculation results in the GUI.
        
        Parameters:
        -----------
        results : dict
            Dictionary containing calculation results
        """
        self.results_display.display_results(results)
        
        # Show success message
        messagebox.showinfo(
            "Calculation Complete",
            f"r Critical = {results['r_critical']:.6f}\n\n"
            f"For significance at α = {results['alpha']}, "
            f"|r| must exceed {results['r_critical']:.4f}"
        )
    
    def analyze_excel_file(self, filepath: str, col1: str, col2: str):
        """
        Analyze Excel file and display correlation results.
        
        Parameters:
        -----------
        filepath : str
            Path to Excel file
        col1, col2 : str
            Column names to correlate
        """
        # Disable analyze button
        self.excel_frame.analyze_button.configure(state="disabled")
        
        # Run analysis in thread
        def run_analysis():
            try:
                # Get parameters
                alpha = float(self.excel_alpha_entry.get())
                test_type = self.excel_test_type_var.get()
                
                # Validate alpha
                is_valid, error_msg = validate_alpha(alpha)
                if not is_valid:
                    self.after(0, lambda: messagebox.showerror("Input Error", error_msg))
                    return
                
                # Analyze file
                success, results, error = analyze_excel_correlation(
                    filepath, col1, col2, alpha, test_type
                )
                
                if not success:
                    self.after(0, lambda: messagebox.showerror("Analysis Error", error))
                    return
                
                # Store results
                self.current_results = results
                
                # Display results
                self.after(0, lambda: self.display_excel_results(results))
                
            except Exception as e:
                self.after(0, lambda: messagebox.showerror(
                    "Analysis Error",
                    f"An error occurred:\n{str(e)}"
                ))
            finally:
                self.after(0, lambda: self.excel_frame.analyze_button.configure(state="normal"))
        
        thread = threading.Thread(target=run_analysis)
        thread.daemon = True
        thread.start()
    
    def display_excel_results(self, results: dict):
        """Display Excel analysis results."""
        self.excel_results_display.display_results(results)
        
        # Show summary message
        sig_text = "SIGNIFICANT" if results['is_significant'] else "NOT SIGNIFICANT"
        messagebox.showinfo(
            "Analysis Complete",
            f"Correlation between:\n"
            f"  {results['column_1']} and {results['column_2']}\n\n"
            f"Pearson's r = {results['r_value']:.6f}\n"
            f"P-value = {results['p_value']:.6f}\n\n"
            f"Result: {sig_text}\n"
            f"(|r| = {abs(results['r_value']):.4f} vs threshold {results['r_critical']:.4f})"
        )
    
    def save_results(self):
        """Save current results to a Word document (.docx) with file dialog."""
        if self.current_results is None:
            messagebox.showwarning(
                "No Results",
                "Please calculate results before saving."
            )
            return
        
        try:
            from tkinter import filedialog
            
            # Create default filename with timestamp
            timestamp = datetime.now().strftime("%m-%d-%Y_%I-%M%p")
            default_filename = f"pearson_r_results_{timestamp}.docx"
            
            # Open file dialog for user to choose location and filename
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[
                    ("Word Documents", "*.docx"),
                    ("All Files", "*.*")
                ],
                initialfile=default_filename,
                title="Save Results As"
            )
            
            # If user cancels, filepath will be empty
            if not filepath:
                return
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading("Pearson's r Critical Value Analysis", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            timestamp_run.italic = True
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Check if Excel analysis
            is_excel = 'column_1' in self.current_results and 'r_value' in self.current_results
            
            if is_excel:
                self.save_excel_results_to_doc(doc)
            else:
                self.save_manual_results_to_doc(doc)
            
            # Save document
            doc.save(filepath)
            
            messagebox.showinfo(
                "Save Successful",
                f"Results saved to:\n{filepath}"
            )
            
        except Exception as e:
            messagebox.showerror(
                "Save Error",
                f"Failed to save results:\n{str(e)}"
            )
    
    def save_manual_results_to_doc(self, doc):
        """Save manual calculation results to document."""
        doc.add_heading("Input Parameters", level=1)
        
        input_table = doc.add_table(rows=4, cols=2)
        input_table.style = 'Light Grid Accent 1'
        
        header_cells = input_table.rows[0].cells
        header_cells[0].text = "Parameter"
        header_cells[1].text = "Value"
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        input_table.rows[1].cells[0].text = "Test Type"
        input_table.rows[1].cells[1].text = self.current_results['test_type'].title()
        
        input_table.rows[2].cells[0].text = "Sample Size (n)"
        input_table.rows[2].cells[1].text = str(self.current_results['sample_size'])
        
        input_table.rows[3].cells[0].text = "Significance Level (α)"
        input_table.rows[3].cells[1].text = str(self.current_results['alpha'])
        
        doc.add_paragraph()
        
        doc.add_heading("Calculation Results", level=1)
        
        results_table = doc.add_table(rows=4, cols=2)
        results_table.style = 'Light Grid Accent 1'
        
        results_header = results_table.rows[0].cells
        results_header[0].text = "Statistic"
        results_header[1].text = "Value"
        
        for cell in results_header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        results_table.rows[1].cells[0].text = "Degrees of Freedom (df)"
        results_table.rows[1].cells[1].text = str(self.current_results['degrees_of_freedom'])
        
        results_table.rows[2].cells[0].text = "t Critical"
        results_table.rows[2].cells[1].text = f"{self.current_results['t_critical']:.6f}"
        
        r_row = results_table.rows[3].cells
        r_row[0].text = "r Critical"
        r_row[1].text = f"{self.current_results['r_critical']:.6f}"
        
        for cell in r_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
    
    def save_excel_results_to_doc(self, doc):
        """Save Excel analysis results to document."""
        doc.add_heading("Excel Correlation Analysis", level=1)
        
        # Variables section
        doc.add_heading("Variables Analyzed", level=2)
        var_para = doc.add_paragraph()
        var_para.add_run(f"X Variable: ").bold = True
        var_para.add_run(f"{self.current_results['column_1']}\n")
        var_para.add_run(f"Y Variable: ").bold = True
        var_para.add_run(self.current_results['column_2'])
        
        doc.add_paragraph()
        
        # Correlation results
        doc.add_heading("Correlation Results", level=2)
        
        corr_table = doc.add_table(rows=5, cols=2)
        corr_table.style = 'Light Grid Accent 1'
        
        header = corr_table.rows[0].cells
        header[0].text = "Statistic"
        header[1].text = "Value"
        
        corr_table.rows[1].cells[0].text = "Pearson's r"
        corr_table.rows[1].cells[1].text = f"{self.current_results['r_value']:.6f}"
        
        corr_table.rows[2].cells[0].text = "P-value"
        corr_table.rows[2].cells[1].text = f"{self.current_results['p_value']:.6f}"
        
        corr_table.rows[3].cells[0].text = "Sample Size (n)"
        corr_table.rows[3].cells[1].text = str(self.current_results['sample_size'])
        
        corr_table.rows[4].cells[0].text = "Significance Level (α)"
        corr_table.rows[4].cells[1].text = str(self.current_results['alpha'])
        
        doc.add_paragraph()
        
        # Interpretation
        doc.add_heading("Statistical Significance", level=2)
        
        interp = doc.add_paragraph()
        sig_text = "SIGNIFICANT" if self.current_results['is_significant'] else "NOT SIGNIFICANT"
        interp.add_run(f"Result: {sig_text}\n\n").bold = True
        interp.add_run(self.current_results['significance_interpretation'])


def main():
    """Main entry point for the application."""
    app = PearsonCalculatorApp()
    app.mainloop()


if __name__ == "__main__":
    main()