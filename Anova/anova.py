from scipy.stats import f_oneway
import numpy as np
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

def get_user_input():
    """
    Get all user inputs from console
    """
    print("🚀 ANOVA REPORT GENERATOR")
    print("=" * 50)
    
    # Get title and subtitle
    title = input("Enter report title (e.g., ANOVA ANALYSIS RESULTS): ").strip()
    if not title:
        title = "ANOVA ANALYSIS RESULTS"
    
    subtitle = input("Enter subtitle (e.g., Age Profile vs Strategies in Facilitating): ").strip()
    if not subtitle:
        subtitle = "Age Profile vs Strategies in Facilitating"
    
    filename = input("Enter output filename (without .docx): ").strip()
    if not filename:
        filename = "ANOVA_Results"
    filename += ".docx"
    
    print("\n📊 Enter data for each age group (comma-separated values)")
    print("Example: 3.27, 3.47, 3.53, 3.27, 3.6")
    print("-" * 50)
    
    # Age group labels
    age_labels = [
        "Below 25",
        "25-29", 
        "30-34",
        "35-39",
        "40-44",
        "45-49",
        "50-54",
        "55 above"
    ]
    
    all_groups = []
    
    for i, age_label in enumerate(age_labels, 1):
        while True:
            try:
                user_input = input(f"Group {i} ({age_label}): ").strip()
                
                if not user_input:
                    print("⚠️  Please enter some values")
                    continue
                
                # Convert input to list of floats
                data_list = [float(x.strip()) for x in user_input.split(',')]
                
                if len(data_list) < 2:
                    print("⚠️  Please enter at least 2 values")
                    continue
                
                all_groups.append(data_list)
                print(f"✅ Added {len(data_list)} values")
                break
                
            except ValueError:
                print("❌ Invalid input. Please enter numbers only (e.g., 3.27, 4.0, 2.5)")
            except Exception as e:
                print(f"❌ Error: {e}")
    
    return title, subtitle, filename, all_groups, age_labels

def manual_anova(groups):
    """
    Calculate ANOVA manually to get SS, MS values
    """
    groups = [np.array(group) for group in groups]
    all_data = np.concatenate(groups)
    grand_mean = np.mean(all_data)
    k = len(groups)
    N = len(all_data)
    
    SS_total = np.sum((all_data - grand_mean) ** 2)
    
    SS_between = 0
    for group in groups:
        n_i = len(group)
        group_mean = np.mean(group)
        SS_between += n_i * (group_mean - grand_mean) ** 2
    
    SS_within = SS_total - SS_between
    
    df_between = k - 1
    df_within = N - k
    df_total = N - 1
    
    MS_between = SS_between / df_between
    MS_within = SS_within / df_within
    
    F_calculated = MS_between / MS_within
    
    return SS_between, SS_within, SS_total, df_between, df_within, MS_between, MS_within, F_calculated

def create_anova_report(title, subtitle, filename, all_groups, age_labels):
    """
    Create ANOVA DOCX report with custom title, subtitle, and filename
    """
    # Calculate ANOVA
    SS_between, SS_within, SS_total, df_between, df_within, MS_between, MS_within, F_calculated = manual_anova(all_groups)
    F_statistic, p_value = f_oneway(*all_groups)
    
    # Format p-value
    if p_value < 0.0001:
        p_value_formatted = "< 0.0001"
    else:
        p_value_formatted = f"{p_value:.6f}"
    
    # Create DOCX document
    doc = Document()
    
    # Title
    doc_title = doc.add_heading(title, 0)
    doc_title.alignment = 1  # Center alignment
    
    # Subtitle
    doc_subtitle = doc.add_heading(subtitle, level=2)
    doc_subtitle.alignment = 1
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = 1
    date_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_paragraph()  # Empty line
    
    # Descriptive Statistics Section
    doc.add_heading('Descriptive Statistics', level=2)
    
    # Create descriptive statistics table
    desc_table = doc.add_table(rows=len(all_groups)+2, cols=4)
    desc_table.style = 'Table Grid'
    
    # Header row
    header_cells = desc_table.rows[0].cells
    headers = ['Age Group', 'Mean', 'n', 'Std. Dev']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows for each group
    for i, group in enumerate(all_groups):
        row_cells = desc_table.rows[i+1].cells
        mean = np.mean(group)
        std_dev = np.std(group, ddof=1)
        n = len(group)
        
        row_cells[0].text = age_labels[i]
        row_cells[1].text = f"{mean:.3f}"
        row_cells[2].text = str(n)
        row_cells[3].text = f"{std_dev:.4f}"
    
    # Total row
    total_cells = desc_table.rows[len(all_groups)+1].cells
    all_data = np.concatenate(all_groups)
    overall_mean = np.mean(all_data)
    overall_std = np.std(all_data, ddof=1)
    total_n = len(all_data)
    
    total_cells[0].text = "Total"
    total_cells[1].text = f"{overall_mean:.3f}"
    total_cells[2].text = str(total_n)
    total_cells[3].text = f"{overall_std:.4f}"
    
    doc.add_paragraph()  # Empty line
    
    # ANOVA Table Section
    doc.add_heading('ANOVA Table', level=2)
    
    # Create ANOVA table
    anova_table = doc.add_table(rows=4, cols=6)
    anova_table.style = 'Table Grid'
    
    # Header row
    anova_header = anova_table.rows[0].cells
    anova_headers = ['Source', 'SS', 'df', 'MS', 'F', 'p-value']
    for i, header in enumerate(anova_headers):
        anova_header[i].text = header
        for paragraph in anova_header[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    treatment_cells = anova_table.rows[1].cells
    treatment_cells[0].text = 'Treatment'
    treatment_cells[1].text = f"{SS_between:.4f}"
    treatment_cells[2].text = str(df_between)
    treatment_cells[3].text = f"{MS_between:.4f}"
    treatment_cells[4].text = f"{F_calculated:.2f}"
    treatment_cells[5].text = p_value_formatted
    
    error_cells = anova_table.rows[2].cells
    error_cells[0].text = 'Error'
    error_cells[1].text = f"{SS_within:.4f}"
    error_cells[2].text = str(df_within)
    error_cells[3].text = f"{MS_within:.4f}"
    error_cells[4].text = ''
    error_cells[5].text = ''
    
    total_cells_anova = anova_table.rows[3].cells
    total_cells_anova[0].text = 'Total'
    total_cells_anova[1].text = f"{SS_total:.4f}"
    total_cells_anova[2].text = str(df_between + df_within)
    total_cells_anova[3].text = ''
    total_cells_anova[4].text = ''
    total_cells_anova[5].text = ''
    
    doc.add_paragraph()  # Empty line
    
    # Interpretation Section
    doc.add_heading('Statistical Interpretation', level=2)
    
    # Key statistics
    stats_para = doc.add_paragraph()
    stats_para.add_run('F-statistic: ').bold = True
    stats_para.add_run(f'{F_statistic:.4f}\n')
    
    stats_para.add_run('Degrees of Freedom: ').bold = True
    stats_para.add_run(f'({df_between}, {df_within})\n')
    
    stats_para.add_run('p-value: ').bold = True
    stats_para.add_run(f'{p_value_formatted}\n')
    
    stats_para.add_run('Alpha level: ').bold = True
    stats_para.add_run('0.05\n\n')
    
    # Decision
    decision_para = doc.add_paragraph()
    if p_value < 0.05:
        decision_para.add_run('✅ REJECT H₀ → There is a SIGNIFICANT difference between the group means\n').bold = True
        decision_para.add_run(f'{subtitle}')
    else:
        decision_para.add_run('❌ FAIL TO REJECT H₀ → No significant difference between group means').bold = True
    
    doc.add_paragraph()  # Empty line
    
    # Save the document
    doc.save(filename)
    
    # Return all necessary values for console summary
    return filename, F_calculated, p_value_formatted, SS_between, SS_within, p_value, df_between, df_within, MS_between, MS_within

def main():
    """
    Main function - reusable ANOVA report generator
    """
    while True:
        # Get user input
        title, subtitle, filename, all_groups, age_labels = get_user_input()
        
        # Create report
        output_file, F_value, p_value, SS_between, SS_within, p_val_num, df_between, df_within, MS_between, MS_within = create_anova_report(
            title, subtitle, filename, all_groups, age_labels
        )
        
        # Show results
        print("\n" + "=" * 70)
        print("✅ DOCUMENT SUCCESSFULLY CREATED!")
        print("=" * 70)
        print(f"📄 File name: {output_file}")
        print(f"📁 Location: {os.path.abspath(output_file)}")
        print(f"🏷️  Title: {title}")
        print(f"📝 Subtitle: {subtitle}")
        print("=" * 70)
        
        # Console summary
        print("\nQUICK CONSOLE SUMMARY:")
        print("=" * 60)
        print(f"{'Source':<12} {'SS':<10} {'df':<6} {'MS':<10} {'F':<10} {'p-value':<12}")
        print("-" * 65)
        print(f"{'Treatment':<12} {SS_between:.4f}  {df_between:<6} {MS_between:.4f}   {F_value:.2f}    {p_value:<12}")
        print(f"{'Error':<12} {SS_within:.4f}  {df_within:<6} {MS_within:.4f}")
        print(f"{'Total':<12} {SS_between + SS_within:.4f}  {df_between + df_within:<6}")
        
        if p_val_num < 0.05:
            print(f"\n✅ SIGNIFICANT - {subtitle}")
        else:
            print(f"\n❌ NOT SIGNIFICANT - {subtitle}")
        
        print("-" * 60)
        
        # Ask if user wants to create another report
        another = input("\n🔄 Create another ANOVA report? (y/n): ").strip().lower()
        if another not in ['y', 'yes']:
            print("👋 Thank you for using ANOVA Report Generator!")
            break

# Run the program
if __name__ == "__main__":
    main()