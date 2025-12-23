from scipy.stats import f_oneway
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
from statsmodels.stats.multicomp import pairwise_tukeyhsd

def get_user_input():
    print("🚀 ANOVA REPORT GENERATOR")
    print("=" * 50)
    
    title = input("Enter report title (e.g., ANOVA ANALYSIS RESULTS): ").strip() or "ANOVA ANALYSIS RESULTS"
    subtitle = input("Enter subtitle (e.g., Age Profile vs Strategies in Facilitating): ").strip() or "Age Profile vs Strategies in Facilitating"
    filename = input("Enter output filename (without .docx): ").strip() or "ANOVA_Results"
    filename += ".docx"
    
    print("\n📊 Enter data for each age group (space-separated values)")
    print("Example: 3.27 3.47 3.53 3.27 3.6")
    print("-" * 50)
    
    age_labels = ["Below 25","25-29","30-34","35-39","40-44","45-49","50-54","55 above"]
    all_groups = []
    
    for i, age_label in enumerate(age_labels, 1):
        while True:
            try:
                user_input = input(f"Group {i} ({age_label}): ").strip()
                if not user_input:
                    print("⚠️  Please enter some values"); continue
                data_list = [float(x.strip()) for x in user_input.split()]  # space-separated
                if len(data_list) < 2:
                    print("⚠️  Please enter at least 2 values"); continue
                all_groups.append(data_list)
                print(f"✅ Added {len(data_list)} values"); break
            except ValueError:
                print("❌ Invalid input. Please enter numbers only (e.g., 3.27 4.0 2.5)")
    return title, subtitle, filename, all_groups, age_labels

def manual_anova(groups):
    groups = [np.array(g) for g in groups]
    all_data = np.concatenate(groups)
    grand_mean = np.mean(all_data)
    k = len(groups); N = len(all_data)
    
    SS_total = np.sum((all_data - grand_mean)**2)
    SS_between = sum(len(g)*(np.mean(g)-grand_mean)**2 for g in groups)
    SS_within = SS_total - SS_between
    
    df_between = k-1; df_within = N-k; df_total = N-1
    MS_between = SS_between / df_between
    MS_within = SS_within / df_within
    F_calculated = MS_between / MS_within
    
    return SS_between, SS_within, SS_total, df_between, df_within, MS_between, MS_within, F_calculated

def create_anova_report(title, subtitle, filename, all_groups, age_labels):
    SS_between, SS_within, SS_total, df_between, df_within, MS_between, MS_within, F_calculated = manual_anova(all_groups)
    F_statistic, p_value = f_oneway(*all_groups)
    p_value_formatted = "< 0.0001" if p_value<0.0001 else f"{p_value:.6f}"
    
    doc = Document()
    
    # Title
    doc_title = doc.add_heading(title, 0)
    doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_subtitle = doc.add_heading(subtitle, level=2)
    doc_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # Descriptive Statistics Table
    doc.add_heading('Descriptive Statistics', level=2)
    desc_table = doc.add_table(rows=len(all_groups)+2, cols=4)
    desc_table.style = 'Table Grid'
    headers = ['Age Group','Mean','n','Std. Dev']
    for i,h in enumerate(headers): 
        desc_table.rows[0].cells[i].text = h
        for p in desc_table.rows[0].cells[i].paragraphs:
            for run in p.runs: run.bold=True
    
    for i,g in enumerate(all_groups):
        row = desc_table.rows[i+1].cells
        mean, std, n = np.mean(g), np.std(g, ddof=1), len(g)
        row[0].text = age_labels[i]; row[1].text=f"{mean:.3f}"; row[2].text=str(n); row[3].text=f"{std:.4f}"
    
    total_cells = desc_table.rows[len(all_groups)+1].cells
    all_data = np.concatenate(all_groups)
    total_cells[0].text = "Total"
    total_cells[1].text = f"{np.mean(all_data):.3f}"
    total_cells[2].text = str(len(all_data))
    total_cells[3].text = f"{np.std(all_data, ddof=1):.4f}"
    
    # ANOVA Table
    doc.add_paragraph()
    doc.add_heading('ANOVA Table', level=2)
    anova_table = doc.add_table(rows=4, cols=6)
    anova_table.style='Table Grid'
    anova_headers=['Source','SS','df','MS','F','p-value']
    for i,h in enumerate(anova_headers):
        anova_table.rows[0].cells[i].text = h
        for p in anova_table.rows[0].cells[i].paragraphs:
            for run in p.runs: run.bold=True
    # Fill ANOVA rows
    anova_table.rows[1].cells[0].text='Treatment'; anova_table.rows[1].cells[1].text=f"{SS_between:.4f}"; anova_table.rows[1].cells[2].text=str(df_between); anova_table.rows[1].cells[3].text=f"{MS_between:.4f}"; anova_table.rows[1].cells[4].text=f"{F_calculated:.2f}"; anova_table.rows[1].cells[5].text=p_value_formatted
    anova_table.rows[2].cells[0].text='Error'; anova_table.rows[2].cells[1].text=f"{SS_within:.4f}"; anova_table.rows[2].cells[2].text=str(df_within); anova_table.rows[2].cells[3].text=f"{MS_within:.4f}"
    anova_table.rows[3].cells[0].text='Total'; anova_table.rows[3].cells[1].text=f"{SS_total:.4f}"; anova_table.rows[3].cells[2].text=str(df_between+df_within)
    
    # Statistical Interpretation
    doc.add_paragraph(); doc.add_heading('Statistical Interpretation', level=2)
    stats_para = doc.add_paragraph()
    stats_para.add_run(f'F-statistic: ').bold=True; stats_para.add_run(f'{F_statistic:.4f}\n')
    stats_para.add_run('Degrees of Freedom: ').bold=True; stats_para.add_run(f'({df_between}, {df_within})\n')
    stats_para.add_run('p-value: ').bold=True; stats_para.add_run(f'{p_value_formatted}\n')
    stats_para.add_run('Alpha level: ').bold=True; stats_para.add_run('0.05\n\n')
    
    decision_para = doc.add_paragraph()
    posthoc_table = None
    if p_value<0.05:
        decision_para.add_run('✅ REJECT H₀ → SIGNIFICANT difference\n').bold=True
        decision_para.add_run(f'{subtitle}\n')
        
        # Post Hoc Analysis in DOCX
        doc.add_paragraph(); doc.add_heading('Post Hoc Analysis (Tukey HSD)', level=2)
        all_values = np.concatenate(all_groups)
        labels = np.array([[age_labels[i]]*len(g) for i,g in enumerate(all_groups)]).flatten()
        tukey = pairwise_tukeyhsd(all_values, labels)
        tukey_data = tukey.summary().data
        
        posthoc_table = doc.add_table(rows=len(tukey_data), cols=len(tukey_data[0]))
        posthoc_table.style='Table Grid'
        for i,row in enumerate(tukey_data):
            for j,val in enumerate(row):
                posthoc_table.rows[i].cells[j].text = str(val)
                if i==0:
                    for p in posthoc_table.rows[i].cells[j].paragraphs:
                        for run in p.runs: run.bold=True
    else:
        decision_para.add_run('❌ FAIL TO REJECT H₀ → No significant difference\n').bold=True
    
    # Conclusion Section
    doc.add_paragraph(); doc.add_heading('Conclusion', level=2)
    conclusion_para = doc.add_paragraph()
    if p_value<0.05:
        conclusion_para.add_run(f"The ANOVA results indicate a statistically significant difference between the group means for '{subtitle}'.\n").bold=True
        conclusion_para.add_run("Post Hoc (Tukey HSD) analysis identifies which specific groups differ significantly.\n")
    else:
        conclusion_para.add_run(f"The ANOVA results indicate no statistically significant difference between the group means for '{subtitle}'.\n").bold=True
    
    # Hardcoded folder path
    save_folder = "C:\\Users\\kevin\\scipy\\Anova\\Rafael"
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)

    # Full file path
    filepath = os.path.join(save_folder, filename)

    # Save the document
    doc.save(filepath)

    return filename, F_calculated, p_value_formatted, SS_between, SS_within, p_value, df_between, df_within, MS_between, MS_within, posthoc_table

def main():
    while True:
        title, subtitle, filename, all_groups, age_labels = get_user_input()
        output_file, F_value, p_value_str, SS_between, SS_within, p_val_num, df_between, df_within, MS_between, MS_within, posthoc_table = create_anova_report(title, subtitle, filename, all_groups, age_labels)
        
        print("\n✅ DOCUMENT CREATED:", os.path.abspath(output_file))
        
        # Descriptive Statistics
        print("\nDESCRIPTIVE STATISTICS:")
        print(f"{'Age Group':<12} {'Mean':<8} {'n':<4} {'Std Dev':<8}")
        for i,g in enumerate(all_groups):
            print(f"{age_labels[i]:<12} {np.mean(g):<8.3f} {len(g):<4} {np.std(g, ddof=1):<8.4f}")
        print(f"{'Total':<12} {np.mean(np.concatenate(all_groups)):<8.3f} {len(np.concatenate(all_groups)):<4} {np.std(np.concatenate(all_groups), ddof=1):<8.4f}")
        
        # ANOVA Table
        print("\nANOVA TABLE:")
        print(f"{'Source':<12} {'SS':<10} {'df':<6} {'MS':<10} {'F':<10} {'p-value':<12}")
        print(f"{'Treatment':<12} {SS_between:<10.4f} {df_between:<6} {MS_between:<10.4f} {F_value:<10.2f} {p_value_str:<12}")
        print(f"{'Error':<12} {SS_within:<10.4f} {df_within:<6} {MS_within:<10.4f}")
        print(f"{'Total':<12} {SS_between+SS_within:<10.4f} {df_between+df_within:<6}")
        
        # Post Hoc Analysis
        if p_val_num<0.05:
            print("\n✅ SIGNIFICANT → Performing Post Hoc (Tukey HSD)")
            all_values = np.concatenate(all_groups)
            labels = np.array([[age_labels[i]]*len(g) for i,g in enumerate(all_groups)]).flatten()
            tukey = pairwise_tukeyhsd(all_values, labels)
            tukey_data = tukey.summary().data
            print("\nPOST HOC ANALYSIS (Tukey HSD):")
            header = tukey_data[0]
            rows = tukey_data[1:]
            print(f"{header[0]:<12} {header[1]:<12} {header[2]:<12} {header[3]:<10} {header[4]:<10} {header[5]:<10} {header[6]:<10}")
            for r in rows:
                print(f"{r[0]:<12} {r[1]:<12} {r[2]:<12} {r[3]:<10.4f} {r[4]:<10.4f} {r[5]:<10} {r[6]:<10}")
            print("\nCONCLUSION:")
            print(f"The ANOVA results indicate a statistically significant difference between the group means for '{subtitle}'.")
            print("Post Hoc (Tukey HSD) identifies which specific groups differ significantly.")
        else:
            print("\n❌ NOT SIGNIFICANT → Post Hoc not performed")
            print("\nCONCLUSION:")
            print(f"The ANOVA results indicate no statistically significant difference between the group means for '{subtitle}'.")
        
        another = input("\n🔄 Create another report? (y/n): ").strip().lower()
        if another not in ['y','yes']:
            break

if __name__=="__main__":
    main()
